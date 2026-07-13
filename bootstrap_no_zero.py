import sys
import docx

def bootstrap_template_no_zero(src, dest, old_word, new_word):
    print(f"[*] Bootstrapping {dest} from {src}...")
    try:
        doc = docx.Document(src)
        for p in doc.paragraphs:
            if old_word in p.text:
                p.text = p.text.replace(old_word, new_word)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_word in p.text:
                            p.text = p.text.replace(old_word, new_word)
        doc.save(dest)
        print(f"[+] Successfully created {dest} with text replaced.")
    except Exception as e:
        print(f"[-] Error: {e}")

if __name__ == "__main__":
    bootstrap_template_no_zero('Celsis OOS P1 template.docx', 'USP71 OOS P1 template.docx', 'Celsis', 'USP <71>')
