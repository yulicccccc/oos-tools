import docx
import re

def find_malformed_tags(filepath):
    doc = docx.Document(filepath)
    pattern = re.compile(r'\{\{(.*?)\}\}')
    print("Scanning for {{}} tags...")
    for i, p in enumerate(doc.paragraphs):
        matches = pattern.findall(p.text)
        for m in matches:
            # check if it's a valid variable name (alphanumeric and underscores only)
            if not re.match(r'^\s*[a-zA-Z0-9_]+\s*$', m):
                print(f"Malformed tag in paragraph {i}: {{{{ {m} }}}}")
                
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    matches = pattern.findall(p.text)
                    for m in matches:
                        if not re.match(r'^\s*[a-zA-Z0-9_]+\s*$', m):
                            print(f"Malformed tag in table {t_idx}, row {r_idx}, cell {c_idx}: {{{{ {m} }}}}")

if __name__ == "__main__":
    find_malformed_tags('USP71 OOS P1 template 0.docx')
