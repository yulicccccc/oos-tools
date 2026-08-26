import docx
from docx.oxml import parse_xml
import fitz

def fix_word_templates():
    for name in ["EM OOS P1 template.docx", "EM OOS P1 template 0.docx"]:
        doc = docx.Document(name)
        t0 = doc.tables[0]
        
        # Row 20: Is the integrity of the sample(s) questionable? -> Yes (0), No (0), N/A (1)
        r20 = t0.rows[20]
        for cell in r20.cells:
            tc_xml = cell._tc.xml
            if "FORMCHECKBOX" in tc_xml:
                for p in cell.paragraphs:
                    p_xml = p._p.xml
                    if "<w:checkBox>" in p_xml:
                        parts = p_xml.split("<w:checkBox>")
                        if len(parts) >= 4:
                            p1 = parts[1].replace('w:val="1"', 'w:val="0"')
                            p2 = parts[2].replace('w:val="1"', 'w:val="0"')
                            p3 = parts[3].replace('w:val="0"', 'w:val="1"')
                            new_p_xml = parts[0] + "<w:checkBox>" + p1 + "<w:checkBox>" + p2 + "<w:checkBox>" + p3
                            parent = p._p.getparent()
                            new_p = parse_xml(new_p_xml)
                            parent.replace(p._p, new_p)
                            p._p = new_p

        # Row 21: Are the other test results from this run under investigation? -> Yes (0), No (0), N/A (1)
        r21 = t0.rows[21]
        for cell in r21.cells:
            tc_xml = cell._tc.xml
            if "FORMCHECKBOX" in tc_xml:
                for p in cell.paragraphs:
                    p_xml = p._p.xml
                    if "<w:checkBox>" in p_xml:
                        parts = p_xml.split("<w:checkBox>")
                        if len(parts) >= 4:
                            p1 = parts[1].replace('w:val="1"', 'w:val="0"')
                            p2 = parts[2].replace('w:val="1"', 'w:val="0"')
                            p3 = parts[3].replace('w:val="0"', 'w:val="1"')
                            new_p_xml = parts[0] + "<w:checkBox>" + p1 + "<w:checkBox>" + p2 + "<w:checkBox>" + p3
                            parent = p._p.getparent()
                            new_p = parse_xml(new_p_xml)
                            parent.replace(p._p, new_p)
                            p._p = new_p

        doc.save(name)
        print(f"Fixed sample NA checkboxes in Word template: {name}")

def fix_pdf_template():
    pdf_path = "EM OOS P1 template.pdf"
    doc = fitz.open(pdf_path)
    
    # Check Box31 (No=32, NA=33), Check Box34 (Yes=34, No=35, NA=36), Check Box37 (No=38, NA=39)
    # Set 33, 36, 39 to Yes and 31, 32, 34, 35, 37, 38 to Off
    na_boxes = {33, 36, 39}
    off_boxes = {31, 32, 34, 35, 37, 38}
    
    for page in doc:
        for widget in page.widgets():
            if widget.field_name.startswith("Check Box"):
                try:
                    b_num = int(widget.field_name.replace("Check Box", ""))
                    if b_num in na_boxes:
                        widget.field_value = "Yes"
                        widget.update()
                        print(f"Set {widget.field_name} to Yes (Checked)")
                    elif b_num in off_boxes:
                        widget.field_value = "Off"
                        widget.update()
                        print(f"Set {widget.field_name} to Off (Unchecked)")
                except: pass
                
    doc.save(pdf_path, incremental=True, encryption=fitz.PDF_ENCRYPT_KEEP)
    doc.close()
    print("Fixed PDF template checkboxes for Material Discrepancy & Sample Integrity!")

if __name__ == "__main__":
    fix_word_templates()
    fix_pdf_template()
