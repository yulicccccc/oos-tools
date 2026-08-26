import docx
from docx.oxml import parse_xml

for name in ["EM OOS P1 template.docx", "EM OOS P1 template 0.docx"]:
    doc = docx.Document(name)
    t0 = doc.tables[0]
    
    # Row 29: Was the instrument set up correctly? -> 'Please see below'
    r29 = t0.rows[29]
    for i in range(6, 11):
        r29.cells[i].text = "Please see below"
    r29.cells[11].text = "Please see below"
    
    # Row 30: Is the clean room facility certified? -> {{ cr_display }} / {{ cr_exp }}
    r30 = t0.rows[30]
    for i in range(6, 11):
        r30.cells[i].text = "{{ cr_display }}"
    r30.cells[11].text = "{{ cr_exp }}"
    
    # Row 31: System Performance Check (if applicable)? -> N/A
    r31 = t0.rows[31]
    # Update checkboxes: Yes=0, No=0, N/A=1
    for cell in r31.cells:
        tc_xml = cell._tc.xml
        if "FORMCHECKBOX" in tc_xml:
            for p in cell.paragraphs:
                p_xml = p._p.xml
                if "<w:checkBox>" in p_xml:
                    parts = p_xml.split("<w:checkBox>")
                    if len(parts) >= 4:
                        p1 = parts[1].replace('w:val="1"', 'w:val="0"')
                        p2 = parts[2].replace('w:val="1"', 'w:val="0"')
                        p3 = parts[3].replace('w:val="0"', 'w:val="1"')
                        new_p_xml = parts[0] + "<w:checkBox>" + p1 + "<w:checkBox>" + p2 + "<w:checkBox>" + p3
                        parent = p._p.getparent()
                        new_p = parse_xml(new_p_xml)
                        parent.replace(p._p, new_p)
                        p._p = new_p
                        
    for i in range(6, 11):
        r31.cells[i].text = "N/A"
    r31.cells[11].text = "N/A"
    
    doc.save(name)
    print(f"Updated Rows 29, 30, 31 in {name} successfully!")
