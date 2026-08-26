import docx
import re

doc = docx.Document('EM OOS P1 template.docx')
t0 = doc.tables[0]
for r_idx, r in enumerate(t0.rows):
    seen = set()
    for c in r.cells:
        if c._tc in seen: continue
        seen.add(c._tc)
        for p in c.paragraphs:
            p_xml = p._p.xml
            if '<w:checkBox>' in p_xml:
                # find all default values
                defaults = re.findall(r'<w:default w:val="(\d+)"/>', p_xml)
                print(f"Row {r_idx:2d} | text: {p.text.strip()[:40]:40s} | defaults: {defaults}")
