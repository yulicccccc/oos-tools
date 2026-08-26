import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=60, right=60):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def build_tables_for_em():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # --- Table 1 ---
    p1 = doc.add_paragraph()
    r1 = p1.add_run('Table 1: Read Dates & Incubation Observation')
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(9.5)

    t1 = doc.add_table(rows=2, cols=8)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)

    headers_t1 = [
        'ETX Submission ID',
        'Set-up Analyst & Date',
        'Sampling Site & Location',
        'Plate Reading Analyst (≥ 48H)',
        'CFUs Observed after 48 Hour Incubation at 30–35°C (E001031)',
        'Plate Reading Analyst (NLT 5 days)',
        'CFUs Observed after NLT 5-day Incubation at 20–25°C (E001034)',
        'Microbial Identification'
    ]

    for col_idx, text in enumerate(headers_t1):
        cell = t1.cell(0, col_idx)
        cell.text = text
        set_cell_background(cell, 'D9D9D9')
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(8)

    vals_t1 = [
        '{{ etx_id }}',
        '{{ setup_analyst_initial }} {{ test_date }}',
        '{{ sampling_location }}',
        '{{ reader_48h }}',
        '{{ cfu_obs_48h }}',
        '{{ reader_5d }}',
        '{{ cfu_obs_5d }}',
        '{{ microbial_id }}'
    ]

    for col_idx, text in enumerate(vals_t1):
        cell = t1.cell(1, col_idx)
        cell.text = text
        set_cell_margins(cell, top=50, bottom=50, left=50, right=50)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 7 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(8)

    doc.add_paragraph() # space

    # --- Table 2 ---
    p2 = doc.add_paragraph()
    r2 = p2.add_run('Table 2: Environmental Monitoring for Analyst & Cleanroom')
    r2.bold = True
    r2.font.name = 'Arial'
    r2.font.size = Pt(9.5)

    headers_t2 = [
        'Environmental Monitoring (EM) Sampling Site',
        'Frequency',
        'Date (DDMMMYYYY)',
        'Analyst (Initials)',
        'Day /Week(s)',
        'Observation',
        'Environmental Monitoring Plate ETX ID',
        'Microbial ID',
        'Notes'
    ]

    rows_data_t2 = [
        ('HEADER_SECTION', 'Personnel EM Bracketing'),
        ('DATA', ['Personal (Left Touch and Right Touch)', 'Daily', '{{ before_date }}', '{{ analyst_initial }}', 'Date Before Testing', '{{ pers_obs_before }}', '{{ pers_etx_before }}', '{{ pers_id_before }}', 'None']),
        ('DATA', ['Personal (Left Touch and Right Touch)', 'Daily', '{{ test_date }}', '{{ analyst_initial }}', 'Date of Testing', '{{ pers_obs_during }}', '{{ pers_etx_during }}', '{{ pers_id_during }}', 'None']),
        ('DATA', ['Personal (Left Touch and Right Touch)', 'Daily', '{{ after_date }}', '{{ analyst_initial }}', 'Date After Testing', '{{ pers_obs_after }}', '{{ pers_etx_after }}', '{{ pers_id_after }}', 'None']),
        ('HEADER_SECTION', 'Biological Safety Cabinet EM Bracketing Biological Safety Cabinet (BSC)'),
        ('DATA', ['Surface Sampling of ISO 5 (4 locations)', 'Daily', '{{ before_date }}', '{{ bsc_surf_analyst_before }}', 'Date Before Testing', '{{ bsc_surf_obs_before }}', '{{ bsc_surf_etx_before }}', '{{ bsc_surf_id_before }}', 'None']),
        ('DATA', ['Surface Sampling of ISO 5 (4 locations)', 'Daily', '{{ test_date }}', '{{ bsc_surf_analyst_during }}', 'Date of Testing', '{{ bsc_surf_obs_during }}', '{{ bsc_surf_etx_during }}', '{{ bsc_surf_id_during }}', 'None']),
        ('DATA', ['Surface Sampling of ISO 5 (4 locations)', 'Daily', '{{ after_date }}', '{{ bsc_surf_analyst_after }}', 'Date After Testing', '{{ bsc_surf_obs_after }}', '{{ bsc_surf_etx_after }}', '{{ bsc_surf_id_after }}', 'None']),
        ('DATA', ['Settling Sampling of ISO 5 (2 locations)', 'Daily', '{{ before_date }}', '{{ bsc_sett_analyst_before }}', 'Date Before Testing', '{{ bsc_sett_obs_before }}', '{{ bsc_sett_etx_before }}', '{{ bsc_sett_id_before }}', 'None']),
        ('DATA', ['Settling Sampling of ISO 5 (2 locations)', 'Daily', '{{ test_date }}', '{{ bsc_sett_analyst_during }}', 'Date of Testing', '{{ bsc_sett_obs_during }}', '{{ bsc_sett_etx_during }}', '{{ bsc_sett_id_during }}', 'None']),
        ('DATA', ['Settling Sampling of ISO 5 (2 locations)', 'Daily', '{{ after_date }}', '{{ bsc_sett_analyst_after }}', 'Date After Testing', '{{ bsc_sett_obs_after }}', '{{ bsc_sett_etx_after }}', '{{ bsc_sett_id_after }}', 'None']),
        ('HEADER_SECTION', 'Weekly Active Air Sampling Bracketing'),
        ('DATA', ['Active Air Sampling of Cleanrooms', 'Weekly', '{{ date_of_weekly_air }}', '{{ weekly_air_analyst }}', 'Week (On or After Testing Date)', '{{ air_obs }}', '{{ air_etx }}', '{{ air_id }}', 'None']),
        ('HEADER_SECTION', 'Surface Sampling of Anteroom and Cleanroom Bracketing'),
        ('DATA', ['Surface Sampling of Cleanrooms', 'Weekly', '{{ date_of_weekly_surf }}', '{{ weekly_surf_analyst }}', 'Week (On or After Testing Date)', '{{ room_surf_obs }}', '{{ room_surf_etx }}', '{{ room_surf_id }}', 'None'])
    ]

    t2 = doc.add_table(rows=len(rows_data_t2) + 1, cols=9)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)

    for col_idx, text in enumerate(headers_t2):
        cell = t2.cell(0, col_idx)
        cell.text = text
        set_cell_background(cell, 'D9D9D9')
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Arial'
                r.font.size = Pt(8)

    for row_idx, item in enumerate(rows_data_t2, start=1):
        kind, data = item
        if kind == 'HEADER_SECTION':
            first_cell = t2.cell(row_idx, 0)
            for c in range(1, 9):
                first_cell.merge(t2.cell(row_idx, c))
            first_cell.text = data
            set_cell_background(first_cell, 'F2F2F2')
            set_cell_margins(first_cell, top=40, bottom=40, left=60, right=60)
            for p in first_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True
                    r.font.name = 'Arial'
                    r.font.size = Pt(8)
        else:
            for col_idx, val in enumerate(data):
                cell = t2.cell(row_idx, col_idx)
                cell.text = val
                set_cell_margins(cell, top=40, bottom=40, left=50, right=50)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [1,2,3,4,8] else WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(8)

    target_path = 'tables for em.docx'
    doc.save(target_path)
    print(f'Successfully saved {target_path}')

if __name__ == '__main__':
    build_tables_for_em()
