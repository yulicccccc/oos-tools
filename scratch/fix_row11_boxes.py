import docx
from docx.oxml import parse_xml
import pypdf
import fitz

def fix_word_templates():
    for name in ["EM OOS P1 template.docx", "EM OOS P1 template 0.docx"]:
        doc = docx.Document(name)
        t0 = doc.tables[0]
        
        # Row 11: Correct samples analyzed? -> Yes (0), No (0), N/A (1)
        r11 = t0.rows[11]
        for cell in r11.cells:
            tc_xml = cell._tc.xml
            if "FORMCHECKBOX" in tc_xml and "Yes" in cell.text:
                # In this cell, there are 3 checkboxes: Yes, No, N/A
                # We want first checkbox default=0, second default=0, third default=1
                # Let's inspect paragraphs
                for p in cell.paragraphs:
                    p_xml = p._p.xml
                    if "<w:checkBox>" in p_xml:
                        # Replace the 3 checkboxes
                        # First: <w:checkBox><w:sizeAuto/><w:default w:val="1"/></w:checkBox> -> val="0"
                        # Third: <w:checkBox><w:sizeAuto/><w:default w:val="0"/></w:checkBox> -> val="1"
                        parts = p_xml.split("<w:checkBox>")
                        if len(parts) >= 4: # parts[0], parts[1]=first cb, parts[2]=second cb, parts[3]=third cb
                            p1 = parts[1].replace('w:val="1"', 'w:val="0"')
                            p2 = parts[2].replace('w:val="1"', 'w:val="0"')
                            p3 = parts[3].replace('w:val="0"', 'w:val="1"')
                            new_p_xml = parts[0] + "<w:checkBox>" + p1 + "<w:checkBox>" + p2 + "<w:checkBox>" + p3
                            parent = p._p.getparent()
                            new_p = parse_xml(new_p_xml)
                            parent.replace(p._p, new_p)
                            p._p = new_p
        doc.save(name)
        print(f"Fixed Row 11 checkboxes in Word template: {name}")

def fix_pdf_template():
    pdf_path = "EM OOS P1 template.pdf"
    doc = fitz.open(pdf_path)
    
    # Exact checked boxes from G-drive production EM reports
    checked_boxes = {4, 9, 10, 13, 16, 19, 24, 27, 28, 32, 34, 38, 42, 43, 48, 51, 52, 55, 60, 63, 66, 69, 72, 73, 78, 79, 87}
    
    for page in doc:
        for widget in page.widgets():
            if widget.field_name.startswith("Check Box"):
                try:
                    b_num = int(widget.field_name.replace("Check Box", ""))
                    if b_num in checked_boxes:
                        widget.field_value = "Yes"
                    else:
                        widget.field_value = "Off"
                    widget.update()
                except: pass
                
    doc.save(pdf_path, incremental=True, encryption=fitz.PDF_ENCRYPT_KEEP)
    doc.close()
    print("Fixed PDF template checkboxes to 100% match production G-drive standard!")

if __name__ == "__main__":
    fix_word_templates()
    fix_pdf_template()
