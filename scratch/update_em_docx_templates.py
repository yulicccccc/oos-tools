import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def update_templates():
    # 1. Load tables for em.docx to get table definitions
    em_tables_doc = docx.Document('tables for em.docx')
    t1_source = em_tables_doc.tables[0]
    t2_source = em_tables_doc.tables[1]

    for template_name in ['EM OOS P1 template.docx', 'EM OOS P1 template 0.docx']:
        doc = docx.Document(template_name)
        t0 = doc.tables[0]
        
        # Section A Row 2: Test Name
        # In Row 2, cell contains "Test Name: Scan RDI Sterility Test" -> change to "Environmental Monitoring"
        for r_idx in range(len(t0.rows)):
            row = t0.rows[r_idx]
            for cell in row.cells:
                if 'Scan RDI Sterility Test' in cell.text:
                    cell.text = cell.text.replace('Scan RDI Sterility Test', 'Environmental Monitoring')
                if '2.600.023' in cell.text:
                    cell.text = cell.text.replace('2.600.023 (10) 2.700.004 (03)', '2.600.002')
                    cell.text = cell.text.replace('2.600.023', '2.600.002')
                if '07Apr25 02Jan23' in cell.text:
                    cell.text = cell.text.replace('07Apr25 02Jan23', '05 Aug 2025')
                if '10 Rev 03' in cell.text:
                    cell.text = cell.text.replace('10 Rev 03', '15')
                if '0 CFU / 100 mL' in cell.text:
                    cell.text = cell.text.replace('0 CFU / 100 mL', '{{ action_level }}')
                if '{{smart_incident_opening}}' in cell.text or 'The sample was processed' in cell.text:
                    cell.text = '{{ incident_description }}'
                if '{{smart_scan_id}}' in cell.text:
                    cell.text = cell.text.replace('{{smart_scan_id}}', '{{ equipment_summary }}')
                if '{{smart_cr_id}}' in cell.text:
                    cell.text = cell.text.replace('{{smart_cr_id}}', '{{ cr_id }}')

        # Remove old trailing tables (T1, T2, T3) and paragraphs
        # Keep only t0 in doc.tables
        while len(doc.tables) > 1:
            t_to_remove = doc.tables[-1]
            t_to_remove._tbl.getparent().remove(t_to_remove._tbl)

        # Clean up trailing paragraphs after t0
        # Re-add Table 1 and Table 2 paragraphs and tables
        p_t1_title = doc.add_paragraph()
        r1 = p_t1_title.add_run('Table 1: Read Dates & Incubation Observation')
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(9.5)

        # Clone t1
        doc.element.body.append(parse_xml(t1_source._tbl.xml))

        p_space = doc.add_paragraph()

        p_t2_title = doc.add_paragraph()
        r2 = p_t2_title.add_run('Table 2: Environmental Monitoring for Analyst & Cleanroom')
        r2.bold = True
        r2.font.name = 'Arial'
        r2.font.size = Pt(9.5)

        # Clone t2
        doc.element.body.append(parse_xml(t2_source._tbl.xml))

        doc.save(template_name)
        print(f'Successfully updated {template_name}!')

if __name__ == '__main__':
    update_templates()
