import sys
import docx
from docxtpl import DocxTemplate

def bootstrap_template(src, dest, old_word, new_word):
    print(f"[*] Bootstrapping {dest} from {src}...")
    try:
        # Load the docx
        doc = docx.Document(src)
        
        # Replace in paragraphs
        for p in doc.paragraphs:
            if old_word in p.text:
                p.text = p.text.replace(old_word, new_word)
                
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_word in p.text:
                            p.text = p.text.replace(old_word, new_word)
        
        # Save new document
        doc.save(dest)
        print(f"[+] Successfully created {dest} with text replaced.")
        
        # Extract variables using docxtpl
        tpl = DocxTemplate(dest)
        vars = tpl.get_undeclared_template_variables()
        print("\n" + "="*40)
        print("FOUND THE FOLLOWING {{VARIABLES}}:")
        print("="*40)
        for v in sorted(vars):
            print(f" - {v}")
        print("="*40)
            
    except Exception as e:
        print(f"[-] Error: {e}")

if __name__ == "__main__":
    bootstrap_template('Celsis OOS P1 template 0.docx', 'USP71 OOS P1 template 0.docx', 'Celsis', 'USP <71>')
