# filename: em_logic.py
import streamlit as st
import os
import re
import json
import io
import sys
from datetime import datetime, timedelta

# ReportLab for dynamic Page 7 Table Generation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- 1. Central Utilities ---
try:
    from utils import get_room_logic as u_grl, get_full_name, ordinal, num_to_words, get_cleanroom_narrative, get_monthly_cleaning_date
except ImportError:
    def u_grl(i): return "Unknown", "000", "", "Unknown"
    def get_full_name(i): return i
    def ordinal(n): return str(n)
    def num_to_words(n): return str(n)
    def get_cleanroom_narrative(s, r=None, a="", v=""): return ""
    def get_monthly_cleaning_date(p): return "26-Apr-2026"

# --- 2. CONFIG & KEYS ---
FIELD_KEYS = [
    "oos_id", "sample_name", "test_date", "sampling_type", "bsc_id",
    "analyst_name", "analyst_initial", "reader_name", "reader_initial",
    "action_level", "cfu_count", "event_number", "manual_org", "monthly_cleaning_date",
    "cleaner_name", "writer_name", "manager_name", "test_method",
    # Media Plate / Reagents
    "plate_media_type", "media_plate_lot", "media_plate_exp",
    # Bracketing fields
    "before_date", "after_date",
    "pers_obs_before", "pers_obs_during", "pers_obs_after",
    "bsc_surf_obs_before", "bsc_surf_obs_during", "bsc_surf_obs_after",
    "bsc_sett_obs_before", "bsc_sett_obs_during", "bsc_sett_obs_after",
    "date_of_weekly_air", "weekly_air_analyst", "air_obs", "air_etx", "air_id",
    "date_of_weekly_surf", "weekly_surf_analyst", "room_surf_obs", "room_surf_etx", "room_surf_id"
]

def auto_fill_name(initial_key, name_key):
    initial = st.session_state.get(initial_key, "")
    current_name = st.session_state.get(name_key, "")
    if initial:
        calculated_name = get_full_name(initial)
        if calculated_name and not current_name:
            st.session_state[name_key] = calculated_name

def validate_inputs():
    errors, warnings = [], []
    reqs = {
        "OOS Number": "oos_id", "Sample / Plate Name": "sample_name", 
        "Test Date": "test_date", "Setup Analyst Name": "analyst_name",
        "BSC / Cleanroom ID": "bsc_id"
    }
    for label, key in reqs.items():
        if not st.session_state.get(key, "").strip(): warnings.append(label)
    date_val = st.session_state.get("test_date", "").strip()
    if date_val:
        try:
            clean_d = re.sub(r'[\s\-]', '', date_val)
            if len(clean_d) >= 9:
                datetime.strptime(clean_d, "%d%b%Y")
            else:
                datetime.strptime(clean_d, "%d%b%y")
        except ValueError:
            errors.append(f"❌ Date Error: '{date_val}' invalid. Use DDMMMYY (e.g. 17Feb26).")
    return errors, warnings

def clean_filename(text): 
    return re.sub(r'[\\/*?:"<>|]', '_', str(text)).strip() if text else ""

def format_date_std(date_str):
    """Converts diverse date strings to DD-Mon-YYYY format"""
    if not date_str:
        return ""
    clean_d = re.sub(r'[\s\-]', '', str(date_str).strip())
    for fmt in ["%d%b%Y", "%d%b%y", "%Y%m%d", "%m/%d/%Y", "%d/%m/%Y"]:
        try:
            dt = datetime.strptime(clean_d, fmt)
            return dt.strftime("%d-%b-%Y")
        except ValueError:
            continue
    return str(date_str)

def parse_em_text(text):
    """Smart Paste parser for EM email & notification text"""
    data = {}
    if not text or not text.strip():
        return data
        
    # 1. OOS ID
    oos_match = re.search(r"OOS[-\s]*(\d+)", text, re.IGNORECASE)
    if oos_match:
        data["oos_id"] = f"OOS-{oos_match.group(1).strip()}"
        
    # 2. ETX / Event ID
    etx_match = re.search(r"(ETX-\d{6}-\d{4})", text, re.IGNORECASE)
    if etx_match:
        data["event_number"] = etx_match.group(1).strip()
        
    # 3. Plate / Sample Name
    plate_match = re.search(r"((?:Scan|Sterility|EM)[^\t\r\n]+)", text)
    if plate_match:
        p_name = plate_match.group(1).strip()
        data["sample_name"] = p_name
        
        # Extract date from plate name
        date_in_plate = re.search(r"(\d{1,2}\s*[A-Za-z]{3}\s*\d{2,4})", p_name)
        if date_in_plate:
            raw_d = date_in_plate.group(1).replace(" ", "")
            try:
                if len(raw_d) >= 9:
                    d_obj = datetime.strptime(raw_d, "%d%b%Y")
                else:
                    d_obj = datetime.strptime(raw_d, "%d%b%y")
                data["test_date"] = d_obj.strftime("%d%b%y")
            except: pass
            
        # Extract BSC / Equipment ID
        bsc_match = re.search(r"(?:BSC|E00)?(\d{4})", p_name, re.IGNORECASE)
        if bsc_match:
            data["bsc_id"] = f"BSC E00{bsc_match.group(1)}"
            
        # Extract Setup Analyst Initial
        analyst_match = re.search(r"(?:ScanC/O|ScanCO|Scan|Sterility|EM)\s+([A-Z]{2,3})\b", p_name, re.IGNORECASE)
        if analyst_match:
            init = analyst_match.group(1).upper()
            full_n = get_full_name(init)
            if full_n and full_n != init:
                data["analyst_name"] = full_n
                data["analyst_initial"] = init

        # Infer sampling type
        p_lower = p_name.lower()
        if "sett" in p_lower:
            data["sampling_type"] = "Settling Sampling"
        elif "c/o" in p_lower or "changeover" in p_lower:
            data["sampling_type"] = "Surface Sampling (Changeover)"
        elif any(s in p_lower for s in ["s1", "s2", "s3", "s4", "surf"]):
            data["sampling_type"] = "Surface Sampling"
        elif "glove" in p_lower or "pers" in p_lower:
            data["sampling_type"] = "Personnel Sampling (Glove)"
        elif "cart" in p_lower or "floor" in p_lower or "room" in p_lower or "air" in p_lower:
            data["sampling_type"] = "Weekly Cleanroom Sampling"

    # 4. CFU Count
    cfu_match = re.search(r"(?:Total CFU Count on Plate|CFU Count|CFU)\s*[:\n\r]*\s*(\d+)", text, re.IGNORECASE)
    if cfu_match:
        data["cfu_count"] = cfu_match.group(1).strip()
        
    # 5. Colony Description & Organism Identification
    org_match = re.search(r"(?:Microbial Identification|Colony Description|Organism)\s*(?:\(Optional\))?\s*[:\n\r]*\s*([^\n\r]+)", text, re.IGNORECASE)
    if org_match and org_match.group(1).strip().upper() not in ["N/A", "NONE", ""]:
        data["manual_org"] = org_match.group(1).strip()

    # 6. Reagent / Plate Media Lot & Exp
    lot_match = re.search(r"(?:TSA Lot|Contact Plate Lot|Plate Lot|Lot\s*#?|Media Lot)\s*[:\s]*(\d{7,10})", text, re.IGNORECASE)
    if not lot_match:
        lot_match = re.search(r"\b(1011\d{6})\b", text)
    if lot_match:
        data["media_plate_lot"] = lot_match.group(1).strip()
        
    exp_match = re.search(r"(?:Exp|Expiry|Expiration)\s*[:\s]*(\d{1,2}\s*[A-Za-z]{3}\s*\d{2,4})", text, re.IGNORECASE)
    if not exp_match:
        exp_match = re.search(r"\b(\d{1,2}[A-Za-z]{3}\d{2,4})\b", text)
    if exp_match:
        raw_exp = exp_match.group(1).replace(" ", "").upper()
        data["media_plate_exp"] = raw_exp

    # RS Approved Standard Defaults
    data["reader_name"] = "Maraya Chukwumerije and Simin Mohammad"
    data["writer_name"] = "Qiyue Chen"
    data["writer_initial"] = "QYC"
    data["manager_name"] = "Kathan Parikh"
    data["manager_notified"] = "Kathan Parikh"
    data["manager_signer"] = "Robin Seymour"
    data["cleaner_name"] = "Rey Estrada"

    return data

def parse_em_docx_table(docx_input):
    """
    Parses a single EM summary docx file (like 'EM table OOS-261187 11MAY2026.docx')
    and extracts all investigation details, bracketing data, incident notes, and cleanroom mapping.
    """
    if isinstance(docx_input, str):
        d = docx.Document(docx_input)
        src_text = docx_input
    else:
        d = docx.Document(docx_input)
        src_text = getattr(docx_input, 'name', '')

    data = {}
    oos_match = re.search(r'OOS-(\d+)', src_text, re.IGNORECASE)
    if oos_match:
        data['oos_id'] = f"OOS-{oos_match.group(1)}"

    for table in d.tables:
        current_section = ''
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            unique_cells = []
            for c in cells:
                if not unique_cells or c != unique_cells[-1]:
                    unique_cells.append(c)
            
            first_cell = unique_cells[0] if unique_cells else ''
            
            if 'Personnel EM Bracketing' in first_cell:
                current_section = 'pers'
                continue
            elif 'Biological Safety Cabinet EM Bracketing' in first_cell:
                current_section = 'bsc'
                bsc_num_m = re.search(r'\(BSC\)\s*(\d+)', first_cell)
                if bsc_num_m:
                    data['bsc_num'] = bsc_num_m.group(1)
                    data['bsc_id'] = f"BSC E00{bsc_num_m.group(1)}"
                continue
            elif 'Weekly Active Air Sampling Bracketing' in first_cell:
                current_section = 'weekly_air'
                suite_m = re.search(r'Bracketing\s*(\d+)', first_cell)
                if suite_m:
                    data['suite_num'] = suite_m.group(1)
                continue
            elif 'Surface Sampling of Anteroom and Cleanroom Bracketing' in first_cell:
                current_section = 'weekly_surf'
                continue
            elif 'Environmental Monitoring (EM) Sampling Site' in first_cell:
                continue
                
            if len(unique_cells) >= 6:
                site = unique_cells[0]
                freq = unique_cells[1]
                date_str = unique_cells[2]
                analyst = unique_cells[3]
                day_timing = unique_cells[4]
                obs = unique_cells[5]
                etx = unique_cells[6] if len(unique_cells) > 6 else ''
                org = unique_cells[7] if len(unique_cells) > 7 else ''
                
                # Check for incident row (where CFU > 0 or ETX present)
                if etx.startswith('ETX-') and 'Air' not in obs:
                    data['incident_date'] = date_str
                    data['incident_analyst_init'] = analyst.split(',')[0].strip()
                    data['analyst_initial'] = data['incident_analyst_init']
                    data['analyst_name'] = get_full_name(data['analyst_initial'])
                    data['event_number'] = etx
                    data['cfu_obs'] = obs
                    
                    cfu_m = re.search(r'(\d+)\s*CFU', obs)
                    data['cfu_count'] = cfu_m.group(1) if cfu_m else '1'
                    
                    site_m = re.search(r'on\s*(S\d+)', obs)
                    site_code = site_m.group(1) if site_m else 'S1'
                    data['site_code'] = site_code
                    
                    if 'artifact' in org.lower() or 'embeded' in org.lower() or 'smooth' in org.lower() or org.startswith('N/A'):
                        clean_org = org.replace('N/A', '').strip()
                        data['manual_org'] = f"colony-like artifact ({clean_org})" if clean_org else "colony-like artifact"
                    else:
                        data['manual_org'] = org
                        
                    clean_d = date_str.replace(' ', '')
                    bsc_short = data.get('bsc_num', '1309')
                    data['sample_name'] = f"ScanC/O {data['analyst_initial']} E00{bsc_short} {site_code} {clean_d}"
                    data['test_date'] = clean_d
                    data['sampling_type'] = 'Surface Sampling (Changeover)'

    if 'test_date' in data:
        raw_mc = get_monthly_cleaning_date(data['test_date'])
        data['monthly_cleaning_date'] = format_date_std(raw_mc) if raw_mc else "26 April 2026"
        
    data['reader_name'] = "Maraya Chukwumerije and Sophia Santamaria"
    data['reader_48h'] = "MC"
    data['reader_5d'] = "SAS"
    data['writer_name'] = "Qiyue Chen"
    data['writer_initial'] = "QYC"
    data['manager_notified'] = "Kathan Parikh"
    data['manager_signer'] = "Robin Seymour"
    data['cleaner_name'] = "Tamiru Kotisso and Cuong Du"
    data['plate_media_type'] = "Contact Plate"
    data['media_plate_lot'] = "1011543730"
    data['media_plate_exp'] = "29SEP2026"
    
    return data

def compute_em_dates(test_date_str, etx_id=""):
    """
    Computes standard EM incubation milestones and OOS initiation/incident dates.
    - Test Date: Setup date (e.g., 04-Jun-2026)
    - 48h Read (30-35°C in E001031): +2d (Mon-Wed) or +4d (Thu-Fri)
    - 5-Day Read / Date of Incident / Date Initiated (20-25°C in E001034):
      NLT 5 days later (concluding on business day).
    """
    clean_d = re.sub(r'[\s\-]', '', str(test_date_str).strip())
    d_obj = None
    for fmt in ["%d%b%Y", "%d%b%y", "%Y%m%d"]:
        try:
            d_obj = datetime.strptime(clean_d, fmt)
            break
        except: pass

    dt_etx = None
    if etx_id:
        etx_match = re.search(r'ETX-(\d{2})(\d{2})(\d{2})-\d+', str(etx_id), re.IGNORECASE)
        if etx_match:
            yy, mm, dd = etx_match.groups()
            try:
                dt_etx = datetime.strptime(f"20{yy}{mm}{dd}", "%Y%m%d")
            except: pass

    if d_obj:
        test_d_std = d_obj.strftime("%d-%b-%Y")
        d_start = d_obj.strftime("%d %b %Y")
        d_start_full = d_obj.strftime("%d %B %Y")
        w = d_obj.weekday() # 0=Mon, 1=Tue, 2=Wed, 3=Thu, 4=Fri, 5=Sat, 6=Sun
        
        # 48 hours incubation read date
        if w in [3, 4]: # Thu -> Mon (+4d), Fri -> Tue (+4d)
            d_48h_dt = d_obj + timedelta(days=4)
        else: # Mon -> Wed (+2d), Tue -> Thu (+2d), Wed -> Fri (+2d)
            d_48h_dt = d_obj + timedelta(days=2)
        d_48h = d_48h_dt.strftime("%d %b %Y")
        d_48h_full = d_48h_dt.strftime("%d %B %Y")

        # NLT 5 days incubation read / initiation date
        if dt_etx:
            d_final_dt = dt_etx
        else:
            if w == 3: # Thu setup -> final read 2nd Mon (+11d)
                d_final_dt = d_obj + timedelta(days=11)
            elif w == 4: # Fri setup -> final read 2nd Mon (+10d)
                d_final_dt = d_obj + timedelta(days=10)
            else: # Mon/Tue/Wed setup -> final read same day next week (+7d)
                d_final_dt = d_obj + timedelta(days=7)
        
        d_5d = d_final_dt.strftime("%d %b %Y")
        d_5d_full = d_final_dt.strftime("%d %B %Y")
        date_initiated = d_final_dt.strftime("%d-%b-%Y")
        date_of_incident = date_initiated
        
        before_dt = d_obj - timedelta(days=3 if w == 0 else 1)
        after_dt = d_obj + timedelta(days=3 if w == 4 else 1)
        
        before_d = before_dt.strftime("%d %b %Y")
        before_d_full = before_dt.strftime("%d %B %Y")
        after_d = after_dt.strftime("%d %b %Y")
        after_d_full = after_dt.strftime("%d %B %Y")
    else:
        test_d_std = str(test_date_str)
        d_start = str(test_date_str)
        d_start_full = str(test_date_str)
        d_48h = "06 Jun 2026"
        d_48h_full = "06 June 2026"
        d_5d = "11 Jun 2026"
        d_5d_full = "11 June 2026"
        date_initiated = "11-Jun-2026"
        date_of_incident = "11-Jun-2026"
        before_d = "03 Jun 2026"
        before_d_full = "03 June 2026"
        after_d = "05 Jun 2026"
        after_d_full = "05 June 2026"

    return {
        "test_date_std": test_d_std,
        "d_start": d_start,
        "d_start_full": d_start_full,
        "d_48h": d_48h,
        "d_48h_full": d_48h_full,
        "d_5d": d_5d,
        "d_5d_full": d_5d_full,
        "date_initiated": date_initiated,
        "date_of_incident": date_of_incident,
        "before_d": before_d,
        "before_d_full": before_d_full,
        "after_d": after_d,
        "after_d_full": after_d_full
    }

def get_cleanroom_info(sample_name="", bsc_id=""):
    """
    Infers Cleanroom Suite, Room Number, and Equipment ID:
    - CR115 (BSCs 1313, 1314) -> CR115 (E001737)
    - CR116 (BSCs 1311, 1312) -> CR116 (E001738)
    - CR117 (BSCs 1309, 1310) -> CR117 (E001739)
    - CR114 (BSCs 1316, 1798) -> CR114 (E001736)
    - L-Suite (BSCs 1938, 1317, 1319) -> CR145 (E001979)
    """
    combined = (str(sample_name) + " " + str(bsc_id)).upper()
    
    # Check BSC Number specifically preceded by BSC or E00, avoid 2026/2025/etc.
    bsc_num_match = re.search(r'(?:BSC|E00)\s*(\d{4})', combined)
    if bsc_num_match and not bsc_num_match.group(1).startswith("20"):
        bsc_num = bsc_num_match.group(1)
    else:
        all_bsc = re.findall(r'\b(1309|1310|1311|1312|1313|1314|1316|1798|1938|1317|1319|1988|1937)\b', combined)
        bsc_num = all_bsc[0] if all_bsc else ""
    
    bsc_e_id = f"BSC E00{bsc_num}" if bsc_num else ""
    
    if bsc_num in ["1314", "1313"] or "115" in combined:
        room_num = "115B" if bsc_num == "1314" else ("115A" if bsc_num == "1313" else ("115B" if "115B" in combined else "115A"))
        suite_num = "115"
        cr_suite = "CR115"
        cr_display = "CR115 (E001737)"
        cr_exp = "Dec 2026"
        if not bsc_e_id:
            bsc_e_id = f"Cleanroom Suite {room_num}"
    elif bsc_num in ["1312", "1311"] or "116" in combined:
        room_num = "116B" if bsc_num == "1312" else ("116A" if bsc_num == "1311" else ("116B" if "116B" in combined else "116A"))
        suite_num = "116"
        cr_suite = "CR116"
        cr_display = "CR116 (E001738)"
        cr_exp = "Dec 2026"
        if not bsc_e_id:
            bsc_e_id = f"Cleanroom Suite {room_num}"
    elif bsc_num in ["1310", "1309"] or "117" in combined:
        room_num = "117B" if bsc_num == "1310" else ("117A" if bsc_num == "1309" else ("117B" if "117B" in combined else "117A"))
        suite_num = "117"
        cr_suite = "CR117"
        cr_display = "CR117 (E001739)"
        cr_exp = "Dec 2026"
        if not bsc_e_id:
            bsc_e_id = f"Cleanroom Suite {room_num}"
    elif bsc_num in ["1316", "1798"] or "114" in combined:
        room_num = "114B" if bsc_num == "1316" else ("114A" if bsc_num == "1798" else ("114B" if "114B" in combined else "114A"))
        suite_num = "114"
        cr_suite = "CR114"
        cr_display = "CR114 (E001736)"
        cr_exp = "Dec 2026"
        if not bsc_e_id:
            bsc_e_id = f"Cleanroom Suite {room_num}"
    elif bsc_num in ["1938", "1317", "1319", "1988", "1937"] or "145" in combined or "L-SUITE" in combined:
        room_num = "145" if bsc_num in ["1938", "1317", "1319"] else "144"
        suite_num = "L-Suite"
        cr_suite = "L-Suite"
        cr_display = "CR145 (E001979)"
        cr_exp = "Dec 2026"
        if not bsc_e_id:
            bsc_e_id = f"Cleanroom Suite {room_num}"
    else:
        room_num = "115B"
        suite_num = "115"
        cr_suite = "CR115"
        cr_display = "CR115 (E001737)"
        cr_exp = "Dec 2026"
        if not bsc_e_id:
            bsc_e_id = "BSC E001314"
        
    return {
        "bsc_e_id": bsc_e_id,
        "bsc_num": bsc_num,
        "room_num": room_num,
        "suite_num": suite_num,
        "cr_suite": cr_suite,
        "cr_display": cr_display,
        "cr_exp": cr_exp
    }

# --- 3. NARRATIVE GENERATION LOGIC (RS Approved Gold Standard) ---
def generate_em_narrative():
    """Generates the standardized 3-part Phase I narrative for Environmental Monitoring OOS matching RS approved gold standard"""
    s = st.session_state
    analyst_name = s.get("analyst_name", "Guanchen (David) Li")
    analyst_init = s.get("analyst_initial", "GL")
    reader_name = s.get("reader_name", "Maraya Chukwumerije and Simin Mohammad")
    sampling_type = s.get("sampling_type", "Surface Sampling")
    bsc_id = s.get("bsc_id", "BSC E001314")
    plate_name = s.get("sample_name", "Sterility GL E001314 S1 04JUN2026")
    test_date = s.get("test_date", "04 Jun 2026")
    event_id = s.get("event_number", "ETX-260615-0424")
    cfu_count = s.get("cfu_count", "1")
    org_identified = s.get("manual_org", "colony-like artifact")
    monthly_cleaning_date = s.get("monthly_cleaning_date", "26 April 2026")
    cleaner_name = s.get("cleaner_name", "Rey Estrada")
    test_method = s.get("test_method", "Sterility" if "Sterility" in plate_name else "SCAN RDI")

    # Cleanroom & Equipment Mapping
    cr_info = get_cleanroom_info(plate_name, bsc_id)
    bsc_e_id = cr_info["bsc_e_id"]
    room_num = cr_info["room_num"]
    suite_num = cr_info["suite_num"]
    cr_suite = cr_info["cr_suite"]

    # Compute Incubation & Milestone Dates
    dates = compute_em_dates(test_date, event_id)
    d_start = dates["d_start"]
    d_start_full = dates["d_start_full"]
    d_48h = dates["d_48h"]
    d_48h_full = dates["d_48h_full"]
    d_5d = dates["d_5d"]
    d_5d_full = dates["d_5d_full"]
    before_d_full = dates["before_d_full"]
    after_d_full = dates["after_d_full"]

    is_cleanroom_weekly = any(k in sampling_type.lower() or k in plate_name.lower() for k in ["weekly", "air", "cart", "floor", "cleanroom"])
    is_artifact = any(k in org_identified.lower() for k in ["artifact", "anomaly", "nonviable", "no growth upon subculture", "could not be confirmed"])

    # Determine reader names split
    if " and " in reader_name:
        r_parts = [p.strip() for p in reader_name.split(" and ")]
        reader_1 = r_parts[0]
        reader_2 = r_parts[1]
    elif "," in reader_name:
        r_parts = [p.strip() for p in reader_name.split(",")]
        reader_1 = r_parts[0]
        reader_2 = r_parts[1]
    elif "Maraya" in reader_name and "Simin" in reader_name:
        reader_1, reader_2 = "Maraya Chukwumerije", "Simin Mohammad"
    elif "Maraya" in reader_name and ("Sophia" in reader_name or "SAS" in reader_name):
        reader_1, reader_2 = "Maraya Chukwumerije", "Sophia Santamaria"
    else:
        reader_1, reader_2 = reader_name, reader_name

    # --- 1. Interview & Storage Block (Field 49) ---
    sampling_lower = sampling_type.lower()
    
    if is_artifact:
        cfu_obs_desc = f"one colony-like artifact was observed on {plate_name.split()[3] if len(plate_name.split()) > 3 else 'Surface Plate #1'}"
        artifact_block = (
            f"The observed artifact was submitted for microbial identification under {event_id}. "
            f"However, following transfer or inoculation onto fresh media, no growth was observed. "
            f"Therefore, the observed artifact could not be confirmed as a viable microbial colony or reported as a confirmed colony forming unit (CFU). "
            f"The lack of growth upon subculture indicates that the observation may have been nonviable material, an artifact associated with agar preparation or pouring, or another non-microbial artifact."
        )
        recovery_nature = "recovery"
    else:
        cfu_obs_desc = f"{cfu_count} colony forming unit (CFU) was observed on {plate_name.split()[3] if len(plate_name.split()) > 3 else sampling_type}"
        artifact_block = (
            f"Based on the observations in Table 1, the {sampling_lower} plate recovery was submitted for microbial identification under {event_id}. "
            f"The recovered microorganism was identified as {org_identified}."
        )
        recovery_nature = "organism identified was transient or recurring"

    if is_cleanroom_weekly:
        interview_block = (
            f"The analyst involved in the {sampling_lower} plate setup, {analyst_name} ({analyst_init}), and the analysts involved in "
            f"reading the plate, {reader_name}, were interviewed comprehensively. Their responses are documented throughout this investigation.\n\n"
            f"The EM plates were stored in accordance with the supplier's recommendations, visually inspected before use, and verified to be within their assigned expiration dates. "
            f"All materials and supplies were disinfected in accordance with MICRO-SOP-9, Cleaning and Disinfecting Procedure for Microbiology. "
            f"The functionality of the incubators was verified through review of data generated by the in-house continuous monitoring system.\n\n"
            f"{sampling_type} was performed by {analyst_name} during weekly environmental monitoring in Cleanroom Suite {room_num} on the date of testing ({d_start_full}), "
            f"in accordance with MICRO-SOP-2, Environmental Monitoring of the Cleanroom Facility.\n\n"
            f"The plates were initially incubated at a temperature of 30-35°C in incubator E001031 for a minimum duration of 48 hours, commencing on {d_start_full}. "
            f"Following completion of the minimum 48 hours of incubation on {d_48h_full}, no microbial growth was observed. The plates were subsequently "
            f"incubated for a minimum of 5 days at 20-25°C in incubator E001034, with incubation ending on {d_5d_full}. At completion of the second incubation period, "
            f"{cfu_obs_desc}. The plate was read by {reader_1} after the initial incubation period and by {reader_2} after the second incubation period. "
            f"Please see Table 1 for detailed information on the observations during the respective incubations.\n\n"
            f"{artifact_block}\n\n"
            f"To determine whether the recovery was transient or recurring, personnel-monitoring plates for {analyst_name} and Cleanroom Suite {room_num} "
            f"environmental-monitoring plates were bracketed to include the date before testing ({before_d_full}), the date of testing ({d_start_full}), and the date after testing ({after_d_full}), as detailed in Table 2."
        )
    else:
        interview_block = (
            f"The analyst involved in the {sampling_lower} plate setup, {analyst_name} ({analyst_init}), and the analysts involved in "
            f"reading the plate, {reader_name}, were interviewed comprehensively. Their responses are documented throughout this investigation.\n\n"
            f"The EM plates were stored in accordance with the supplier's recommendations, visually inspected before use, and verified to be within their assigned expiration dates. "
            f"All materials and supplies were disinfected in accordance with MICRO-SOP-9, Cleaning and Disinfecting Procedure for Microbiology. "
            f"The functionality of the incubators was verified through review of data generated by the in-house continuous monitoring system.\n\n"
            f"{sampling_type} was performed by {analyst_name} in ISO 5 {bsc_e_id} located in room {room_num} in suite {cr_suite} on the date of testing ({d_start_full}), "
            f"in accordance with MICRO-SOP-2, Environmental Monitoring of the Cleanroom Facility.\n\n"
            f"The plates were initially incubated at a temperature of 30-35°C in incubator E001031 for a minimum duration of 48 hours, commencing on {d_start_full}. "
            f"Following completion of the minimum 48 hours of incubation on {d_48h_full}, no microbial growth was observed. The plates were subsequently "
            f"incubated for a minimum of 5 days at 20-25°C in incubator E001034, with incubation ending on {d_5d_full}. At completion of the second incubation period, "
            f"{cfu_obs_desc}. The plate was read by {reader_1} after the initial incubation period and by {reader_2} after the second incubation period. "
            f"Please see Table 1 for detailed information on the observations during the respective incubations.\n\n"
            f"{artifact_block}\n\n"
            f"To determine whether the {recovery_nature} was transient or recurring, personnel-monitoring plates for {analyst_name} and ISO 5 {bsc_e_id} "
            f"environmental-monitoring plates were bracketed to include the date before testing ({before_d_full}), the date of testing ({d_start_full}), and the date after testing ({after_d_full}), as detailed in Table 2."
        )

    # --- 2. EM Records Block (Field 50) ---
    if is_cleanroom_weekly:
        records_block = (
            f"Environmental Monitoring Summary:\n"
            f"Weekly surface and active air sampling for {cr_suite} for the previous week and following week of testing showed no microbial growth.\n\n"
            f"However, the {sampling_type} for {cr_suite} for the week of testing, performed on {d_start_full} by analyst {analyst_init}, exhibited {cfu_count} CFUs on {plate_name}, "
            f"which were identified as {org_identified}. Routine monitoring on the date of testing showed no growth across other monitored locations.\n\n"
            f"During the interview with the analyst, they indicated that no obvious abnormalities or deviations in the testing procedure were observed. "
            f"All materials were disinfected prior to testing. Moreover, the cleanroom suites in {cr_suite} were thoroughly cleaned and prepared before initiating testing as per MICRO-SOP-2 and MICRO-SOP-9.\n\n"
            f"Monthly cleaning and disinfection of the cleanroom suite, including the ISO 8 anteroom ({suite_num}), ISO 7 buffer room ({suite_num}A), ISO 7 cleanroom ({suite_num}B), "
            f"and the ISO 5 biosafety cabinets located within Room {suite_num}B, were performed on {monthly_cleaning_date} by Analyst - {cleaner_name} - in accordance with MICRO-SOP-9, "
            f"Cleaning and Disinfecting Procedure for Microbiology. All H2O2 indicators passed, confirming the successful completion and effectiveness of the monthly cleaning and disinfection activities "
            f"within Rooms {suite_num}, {suite_num}A, and {suite_num}B. Additionally, routine cleaning and disinfection were performed before and after the testing activity in accordance with MICRO-SOP-9.\n\n"
            f"It is important to note that no samples processed within that week in {cr_suite} failed {test_method} testing that week.\n\n"
            f"Based on the available evidence, the recovery of {cfu_count} CFU of {org_identified} from {plate_name} in {room_num} on the date of testing ({d_start_full}) appears to be an isolated event. "
            f"This assessment is supported by the absence of microbial recovery from surrounding monitored areas and negative routine monitoring results throughout the bracketing period.\n\n"
            f"The negative settling, personnel, and bracketing surface monitoring results demonstrate that the critical environment remained in a state of control. "
            f"The available data do not support migration, persistence, or recurrence of contamination within {cr_suite}. Collectively, the evidence supports that the recovery was an isolated, "
            f"transient, and non-recurring event, while established cleaning, disinfection, and aseptic controls remained effective."
        )
    else:
        if is_artifact:
            obs_assessment = (
                f"Based on the available evidence, one colony-like artifact was observed on {plate_name} in ISO 5 {bsc_e_id} on the date of testing ({d_start_full}). "
                f"However, the observation could not be confirmed as a viable microbial CFU because no growth was obtained following inoculation onto fresh media. "
                f"Therefore, the observation may represent a nonviable or non-microbial artifact, including a potential artifact associated with agar preparation or the agar-pouring process.\n\n"
                f"If the observed artifact had represented a viable CFU, it would appear to be an isolated event. This assessment is supported by the absence of microbial recovery from "
                f"{analyst_name}'s personnel-monitoring plates on the date before testing ({before_d_full}), the date of testing ({d_start_full}), and the date after testing ({after_d_full}); "
                f"the absence of growth from {bsc_e_id} surface samples collected on the date before testing and the date after testing; and the absence of growth from ISO 5 settling plates throughout the bracketing period."
            )
        else:
            obs_assessment = (
                f"Based on the available evidence, the recovery of {cfu_count} CFU of {org_identified} from {plate_name} in ISO 5 {bsc_e_id} on the date of testing ({d_start_full}) appears to be an isolated event. "
                f"This assessment is supported by the absence of microbial recovery from {analyst_name}'s personnel-monitoring plates on the date before testing ({before_d_full}) and the date of testing ({d_start_full}); "
                f"the absence of growth from {bsc_e_id} surface samples collected on the date before testing and the subsequent available monitoring date after testing ({after_d_full}); "
                f"and the absence of growth from ISO 5 settling plates throughout the bracketing period."
            )

        records_block = (
            f"Environmental Monitoring Summary:\n"
            f"Personnel monitoring plates for {analyst_name}, including left- and right-touch plates, showed no microbial growth on the date before testing ({before_d_full}), "
            f"the date of testing ({d_start_full}), and the date after testing ({after_d_full}).\n\n"
            f"For ISO 5 {bsc_e_id}, daily surface sampling of four locations showed no microbial growth on the date before testing ({before_d_full}). "
            f"On the date of testing ({d_start_full}), {cfu_count} CFU was recovered from {plate_name.split()[3] if len(plate_name.split()) > 3 else 'Surface #1'} associated with {analyst_name}. "
            f"The recovery was documented under {event_id}; microbial identification indicated {org_identified}. Surface sampling performed on the date after testing ({after_d_full}) showed no microbial growth.\n\n"
            f"Settling sampling of ISO 5 {bsc_e_id}, including two locations, showed no microbial growth on the date before testing ({before_d_full}), "
            f"the date of testing ({d_start_full}), and the date after testing ({after_d_full}).\n\n"
            f"Weekly active-air monitoring of Suite {suite_num} conducted during the week before testing and the week of testing showed no microbial growth.\n\n"
            f"Weekly surface monitoring of the anteroom and cleanroom areas associated with Suite {suite_num} showed no microbial growth during the week before testing or the week of testing.\n\n"
            f"During the interview, the analyst indicated that no obvious abnormalities or deviations occurred during the testing process. All materials were disinfected before testing, "
            f"and the relevant cleanroom and ISO 5 BSC were cleaned and prepared before testing in accordance with MICRO-SOP-2 and MICRO-SOP-9.\n\n"
            f"Monthly cleaning and disinfection of the cleanroom suite, including the ISO 8 anteroom ({suite_num}), ISO 7 buffer room ({suite_num}A), ISO 7 cleanroom ({suite_num}B), "
            f"and the ISO 5 biosafety cabinets located within Room {suite_num}B, were performed on {monthly_cleaning_date} by Analyst - {cleaner_name} - in accordance with MICRO-SOP-9, "
            f"Cleaning and Disinfecting Procedure for Microbiology. All H2O2 indicators passed, confirming the successful completion and effectiveness of the monthly cleaning and disinfection activities "
            f"within Rooms {suite_num}, {suite_num}A, and {suite_num}B. Additionally, routine cleaning and disinfection were performed before and after the testing activity in accordance with MICRO-SOP-9.\n\n"
            f"It is also important to note that no samples processed by {analyst_name} in ISO 5 {bsc_e_id} on the date of testing ({d_start}) failed {test_method} testing.\n\n"
            f"{obs_assessment}\n\n"
            f"The negative ISO 5 settling, personnel, and bracketing surface monitoring results demonstrate that the {bsc_e_id} critical environment remained in a state of control. "
            f"The available data do not support migration, persistence, or recurrence of contamination within ISO 5 {bsc_e_id}. Collectively, the evidence supports that the recovery was an isolated, "
            f"transient, and non-recurring event, while established cleaning, disinfection, and aseptic controls remained effective."
        )

    # --- 3. Phase I Summary / Defensive Conclusion (Field 51 - RS Gold Standard) ---
    summary_block = (
        "Accordingly, no systemic environmental control deficiencies were identified, and no additional corrective or preventive actions "
        "are warranted at this time beyond continued routine environmental monitoring and adherence to approved cleaning, disinfection, and aseptic procedures."
    )

    return interview_block, records_block, summary_block

def build_em_context():
    """Builds a complete context dictionary for rendering DOCX and PDF templates"""
    s = st.session_state
    interview_block, records_block, summary_block = generate_em_narrative()

    analyst_name = s.get('analyst_name', 'Guanchen (David) Li')
    analyst_init = s.get('analyst_initial', 'GL')
    reader_name = s.get('reader_name', 'Maraya Chukwumerije and Simin Mohammad')
    sampling_type = s.get('sampling_type', 'Surface Sampling')
    bsc_id = s.get('bsc_id', 'BSC E001314')
    plate_name = s.get('sample_name', 'Sterility GL E001314 S1 04JUN2026')
    test_date = s.get('test_date', '04 Jun 2026')
    oos_id = s.get('oos_id', 'OOS-261401').replace('OOS-', '')
    event_id = s.get('event_number', 'ETX-260615-0424')
    action_level = s.get('action_level', 'Action level: ≥ 1CFU/Plate')
    cfu_count = s.get('cfu_count', '1')
    org_identified = s.get('manual_org', 'colony-like artifact')
    writer_name = s.get('writer_name', 'Dhvanir Kansara')
    manager_name = s.get('manager_name', 'Kathan Parikh')

    # Cleanroom & Equipment Mapping
    cr_info = get_cleanroom_info(plate_name, bsc_id)
    bsc_e_id = cr_info["bsc_e_id"]
    room_num = cr_info["room_num"]
    cr_display = cr_info["cr_display"]
    cr_exp = cr_info["cr_exp"]

    # Compute Incubation & Milestone Dates
    dates = compute_em_dates(test_date, event_id)
    test_d_std = dates["test_date_std"]
    date_initiated = dates["date_initiated"]
    date_of_incident = dates["date_of_incident"]
    before_d = dates["before_d"]
    after_d = dates["after_d"]

    # Build personnel display block for Section A (RS Approved Format)
    personnel_block = f"{analyst_name}\n({sampling_type} Analyst)\n\n{reader_name}\n({sampling_type} Plate Reader)"

    # Signature line
    if writer_name and writer_name.strip() and writer_name.strip() != analyst_name.strip():
        analyst_sig = f"{analyst_name} (Written by: {writer_name})"
    else:
        analyst_sig = analyst_name

    # Media Plate / Reagent Info
    plate_media_type = s.get('plate_media_type', 'TSA Plate' if ('air' in plate_name.lower() or 'sett' in plate_name.lower()) else 'Contact Plate')
    media_lot = s.get('media_plate_lot', '1011834770')
    media_exp = s.get('media_plate_exp', '25 Sep 2026')
    reagent_lot_str = f"{plate_media_type}:\n{media_lot}"
    reagent_exp_str = f"{plate_media_type}:\n{media_exp}"

    ctx = {
        # General & Section A
        "oos_id": oos_id,
        "sample_id": event_id,
        "event_number": event_id,
        "etx_id": event_id,
        "sample_name": plate_name,
        "lot_number": plate_name,
        "dosage_form": "Plate",
        "test_date": test_d_std,
        "d_start": dates["d_start"],
        "d_48h": dates["d_48h"],
        "d_5d": dates["d_5d"],
        "date_initiated": date_initiated,
        "date_of_incident": date_of_incident,
        "analyst_name": analyst_name,
        "analyst_initial": analyst_init,
        "setup_analyst_initial": analyst_init,
        "reader_name": reader_name,
        "analyst_signature": analyst_sig,
        "analyst_personnel_block": personnel_block,
        "smart_personnel_block": personnel_block,
        "bsc_id": bsc_e_id,
        "equipment_summary": "Incubator E001031 and Incubator E001034",
        "cr_display": cr_display,
        "cr_exp": cr_exp,
        "cr_id": cr_display,
        "action_level": action_level,
        "incident_description": "The CFU count for the environmental monitoring plate exceeded the action level.",
        "smart_incident_opening": "The CFU count for the environmental monitoring plate exceeded the action level.",
        "report_header": event_id,
        "client_name": "Eagle Analytical Internal EM",

        # Narratives
        "smart_comment_interview": f"Yes, analysts {analyst_name} and {reader_name} were comprehensively interviewed.",
        "smart_comment_records": f"Yes, Information is available on Eagletrax under {event_id}",
        "smart_comment_samples": "Yes, as per MICRO-SOP-2",
        "smart_comment_storage": "Yes, as per MICRO-SOP-2",
        "narrative_summary": f"{interview_block}\n\n{records_block}\n\n{summary_block}",
        "smart_phase1_summary": f"{records_block}\n\n{summary_block}",
        "smart_phase1_continued": summary_block,
        "smart_phase1_part1": interview_block,
        "smart_phase1_part2": f"{records_block}\n\n{summary_block}",

        # Media Plate / Reagent Info
        "plate_media_type": plate_media_type,
        "media_plate_lot": media_lot,
        "media_plate_exp": media_exp,
        "reagent_lot": reagent_lot_str,
        "reagent_exp": reagent_exp_str,

        # Table 1 Fields
        "sampling_location": f"{sampling_type} Plate ({bsc_e_id})",
        "reader_48h": s.get('reader_48h', 'MC' if ('maraya' in reader_name.lower() or 'mc' in reader_name.lower()) else 'MC'),
        "cfu_obs_48h": f"No microbial growth was observed",
        "reader_5d": s.get('reader_5d', 'SAS' if ('sophia' in reader_name.lower() or 'sas' in reader_name.lower()) else 'SMO'),
        "cfu_obs_5d": f"{cfu_count} CFU on {plate_name}",
        "microbial_id": org_identified,

        # Table 2 Bracketing Fields
        "before_date": before_d,
        "after_date": after_d,
        "pers_obs_before": "No growth", "pers_etx_before": "N/A", "pers_id_before": "N/A",
        "pers_obs_during": "No growth", "pers_etx_during": "N/A", "pers_id_during": "N/A",
        "pers_obs_after": "No growth", "pers_etx_after": "N/A", "pers_id_after": "N/A",

        "bsc_surf_analyst_before": analyst_init, "bsc_surf_obs_before": "No growth", "bsc_surf_etx_before": "N/A", "bsc_surf_id_before": "N/A",
        "bsc_surf_analyst_during": analyst_init, "bsc_surf_obs_during": f"{cfu_count} CFU" if "Surface" in sampling_type else "No growth", "bsc_surf_etx_during": event_id if "Surface" in sampling_type else "N/A", "bsc_surf_id_during": org_identified if "Surface" in sampling_type else "N/A",
        "bsc_surf_analyst_after": analyst_init, "bsc_surf_obs_after": "No growth", "bsc_surf_etx_after": "N/A", "bsc_surf_id_after": "N/A",

        "bsc_sett_analyst_before": analyst_init, "bsc_sett_obs_before": "No growth", "bsc_sett_etx_before": "N/A", "bsc_sett_id_before": "N/A",
        "bsc_sett_analyst_during": analyst_init, "bsc_sett_obs_during": f"{cfu_count} CFU" if "Settling" in sampling_type else "No growth", "bsc_sett_etx_during": event_id if "Settling" in sampling_type else "N/A", "bsc_sett_id_during": org_identified if "Settling" in sampling_type else "N/A",
        "bsc_sett_analyst_after": analyst_init, "bsc_sett_obs_after": "No growth", "bsc_sett_etx_after": "N/A", "bsc_sett_id_after": "N/A",

        "date_of_weekly_air": before_d, "weekly_air_analyst": "Rey Estrada",
        "air_obs": "No growth", "air_etx": "N/A", "air_id": "N/A",

        "date_of_weekly_surf": before_d, "weekly_surf_analyst": "Rey Estrada",
        "room_surf_obs": "No growth", "room_surf_etx": "N/A", "room_surf_id": "N/A",
        
        "writer_name": writer_name,
        "manager_name": manager_name,
        "manager_notified": s.get('manager_notified', 'Kathan Parikh'),
        "manager_signer": s.get('manager_signer', 'Robin Seymour'),
        "section_c_other": s.get('section_c_other', f"N/A {s.get('writer_initial', 'QYC')} {datetime.now().strftime('%d-%b-%Y')}")
    }
    return ctx

# --- 4. DYNAMIC PAGE 7 ATTACHMENT GENERATOR (ReportLab) ---
def generate_em_tables_page_pdf(ctx):
    """Generates vector Page 7 containing Table 1 & Table 2 matching official QA standards"""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TableTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11,
        textColor=colors.black,
        spaceAfter=4
    )
    cell_hdr_style = ParagraphStyle(
        'HdrCell',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=9,
        alignment=TA_CENTER,
        textColor=colors.black
    )
    cell_body_style = ParagraphStyle(
        'BodyCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=6.5,
        leading=8,
        alignment=TA_CENTER
    )
    cell_body_left = ParagraphStyle(
        'BodyCellLeft',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=6.5,
        leading=8,
        alignment=TA_LEFT
    )

    story = []
    
    # --- TABLE 1: Read Dates and Incubation Observation ---
    story.append(Paragraph("<b>Table 1: Read Dates and Incubation Observation</b>", title_style))
    
    t1_headers = [
        Paragraph("<b>Sampling Location</b>", cell_hdr_style),
        Paragraph("<b>Read Date<br/>(30-35°C, NLT 48h)</b>", cell_hdr_style),
        Paragraph("<b>Read<br/>By</b>", cell_hdr_style),
        Paragraph("<b>CFU Count /<br/>Observation</b>", cell_hdr_style),
        Paragraph("<b>Read Date<br/>(20-25°C, NLT 5d)</b>", cell_hdr_style),
        Paragraph("<b>Read<br/>By</b>", cell_hdr_style),
        Paragraph("<b>CFU Count /<br/>Observation</b>", cell_hdr_style),
        Paragraph("<b>Microbial Identification</b>", cell_hdr_style)
    ]
    
    t1_row = [
        Paragraph(ctx['sampling_location'], cell_body_left),
        Paragraph(ctx.get('d_48h', '06 Jun 2026'), cell_body_style),
        Paragraph(ctx.get('reader_48h', 'MC'), cell_body_style),
        Paragraph(ctx.get('cfu_obs_48h', 'No microbial growth was observed'), cell_body_style),
        Paragraph(ctx.get('d_5d', '11 Jun 2026'), cell_body_style),
        Paragraph(ctx.get('reader_5d', 'SMO'), cell_body_style),
        Paragraph(ctx.get('cfu_obs_5d', '1 CFU on Surface Plate #1'), cell_body_style),
        Paragraph(ctx.get('microbial_id', 'colony-like artifact'), cell_body_style)
    ]
    
    t1_table = Table([t1_headers, t1_row], colWidths=[100, 65, 30, 80, 65, 30, 80, 90])
    t1_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(t1_table)
    story.append(Spacer(1, 10))

    # --- TABLE 2: EM Plates Bracketing Table ---
    story.append(Paragraph("<b>Table 2: Environmental Monitoring Plates for Analyst and Cleanroom Bracketing</b>", title_style))
    
    t2_hdr_row1 = [
        Paragraph("<b>Sampling Type</b>", cell_hdr_style),
        Paragraph("<b>Sampling Location</b>", cell_hdr_style),
        Paragraph(f"<b>Date Before Testing<br/>({ctx.get('before_date', '03 Jun 2026')})</b>", cell_hdr_style),
        Paragraph("", cell_hdr_style),
        Paragraph(f"<b>Date of Testing<br/>({ctx.get('test_date', '04 Jun 2026')})</b>", cell_hdr_style),
        Paragraph("", cell_hdr_style),
        Paragraph(f"<b>Date After Testing<br/>({ctx.get('after_date', '05 Jun 2026')})</b>", cell_hdr_style),
        Paragraph("", cell_hdr_style)
    ]
    
    t2_hdr_row2 = [
        Paragraph("", cell_hdr_style),
        Paragraph("", cell_hdr_style),
        Paragraph("<b>Analyst / Result</b>", cell_hdr_style),
        Paragraph("<b>ETX / Identification</b>", cell_hdr_style),
        Paragraph("<b>Analyst / Result</b>", cell_hdr_style),
        Paragraph("<b>ETX / Identification</b>", cell_hdr_style),
        Paragraph("<b>Analyst / Result</b>", cell_hdr_style),
        Paragraph("<b>ETX / Identification</b>", cell_hdr_style)
    ]

    t2_rows = [
        t2_hdr_row1,
        t2_hdr_row2,
        [
            Paragraph("Personnel Monitoring", cell_body_left),
            Paragraph("Glove Touch (Left / Right)", cell_body_left),
            Paragraph(f"{ctx['analyst_initial']}<br/>{ctx['pers_obs_before']}", cell_body_style),
            Paragraph(f"{ctx['pers_etx_before']}<br/>{ctx['pers_id_before']}", cell_body_style),
            Paragraph(f"{ctx['analyst_initial']}<br/>{ctx['pers_obs_during']}", cell_body_style),
            Paragraph(f"{ctx['pers_etx_during']}<br/>{ctx['pers_id_during']}", cell_body_style),
            Paragraph(f"{ctx['analyst_initial']}<br/>{ctx['pers_obs_after']}", cell_body_style),
            Paragraph(f"{ctx['pers_etx_after']}<br/>{ctx['pers_id_after']}", cell_body_style)
        ],
        [
            Paragraph("Surface Sampling", cell_body_left),
            Paragraph(f"ISO 5 {ctx['bsc_id']}", cell_body_left),
            Paragraph(f"{ctx['bsc_surf_analyst_before']}<br/>{ctx['bsc_surf_obs_before']}", cell_body_style),
            Paragraph(f"{ctx['bsc_surf_etx_before']}<br/>{ctx['bsc_surf_id_before']}", cell_body_style),
            Paragraph(f"{ctx['bsc_surf_analyst_during']}<br/>{ctx['bsc_surf_obs_during']}", cell_body_style),
            Paragraph(f"{ctx['bsc_surf_etx_during']}<br/>{ctx['bsc_surf_id_during']}", cell_body_style),
            Paragraph(f"{ctx['bsc_surf_analyst_after']}<br/>{ctx['bsc_surf_obs_after']}", cell_body_style),
            Paragraph(f"{ctx['bsc_surf_etx_after']}<br/>{ctx['bsc_surf_id_after']}", cell_body_style)
        ],
        [
            Paragraph("Settling Sampling", cell_body_left),
            Paragraph(f"ISO 5 {ctx['bsc_id']}", cell_body_left),
            Paragraph(f"{ctx['bsc_sett_analyst_before']}<br/>{ctx['bsc_sett_obs_before']}", cell_body_style),
            Paragraph(f"{ctx['bsc_sett_etx_before']}<br/>{ctx['bsc_sett_id_before']}", cell_body_style),
            Paragraph(f"{ctx['bsc_sett_analyst_during']}<br/>{ctx['bsc_sett_obs_during']}", cell_body_style),
            Paragraph(f"{ctx['bsc_sett_etx_during']}<br/>{ctx['bsc_sett_id_during']}", cell_body_style),
            Paragraph(f"{ctx['bsc_sett_analyst_after']}<br/>{ctx['bsc_sett_obs_after']}", cell_body_style),
            Paragraph(f"{ctx['bsc_sett_etx_after']}<br/>{ctx['bsc_sett_id_after']}", cell_body_style)
        ],
        [
            Paragraph("Weekly Cleanroom Active Air", cell_body_left),
            Paragraph(f"Cleanroom {ctx.get('cr_display', 'CR115')}", cell_body_left),
            Paragraph(f"{ctx['weekly_air_analyst']}<br/>{ctx['air_obs']}", cell_body_style),
            Paragraph(f"{ctx['air_etx']}<br/>{ctx['air_id']}", cell_body_style),
            Paragraph(f"{ctx['weekly_air_analyst']}<br/>{ctx['air_obs']}", cell_body_style),
            Paragraph(f"{ctx['air_etx']}<br/>{ctx['air_id']}", cell_body_style),
            Paragraph(f"{ctx['weekly_air_analyst']}<br/>{ctx['air_obs']}", cell_body_style),
            Paragraph(f"{ctx['air_etx']}<br/>{ctx['air_id']}", cell_body_style)
        ],
        [
            Paragraph("Weekly Cleanroom Surface", cell_body_left),
            Paragraph(f"Anteroom & Buffer ({ctx.get('cr_display', 'CR115')})", cell_body_left),
            Paragraph(f"{ctx['weekly_surf_analyst']}<br/>{ctx['room_surf_obs']}", cell_body_style),
            Paragraph(f"{ctx['room_surf_etx']}<br/>{ctx['room_surf_id']}", cell_body_style),
            Paragraph(f"{ctx['weekly_surf_analyst']}<br/>{ctx['room_surf_obs']}", cell_body_style),
            Paragraph(f"{ctx['room_surf_etx']}<br/>{ctx['room_surf_id']}", cell_body_style),
            Paragraph(f"{ctx['weekly_surf_analyst']}<br/>{ctx['room_surf_obs']}", cell_body_style),
            Paragraph(f"{ctx['room_surf_etx']}<br/>{ctx['room_surf_id']}", cell_body_style)
        ]
    ]

    t2_table = Table(t2_rows, colWidths=[90, 90, 60, 70, 60, 70, 60, 70])
    t2_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 1), colors.HexColor('#E8E8E8')),
        ('SPAN', (0, 0), (0, 1)),
        ('SPAN', (1, 0), (1, 1)),
        ('SPAN', (2, 0), (3, 0)),
        ('SPAN', (4, 0), (5, 0)),
        ('SPAN', (6, 0), (7, 0)),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(t2_table)

    doc.build(story)
    buf.seek(0)
    return buf

# --- 5. REPORT GENERATION ENGINE (DOCX & 7-Page PDF) ---
def generate_em_reports():
    """Generates both the official DOCX and complete 7-Page interactive PDF reports"""
    ctx = build_em_context()
    interview_block, records_block, summary_block = generate_em_narrative()

    docx_buf = io.BytesIO()
    pdf_buf = io.BytesIO()

    # 1. Generate Word Document
    target_docx = "EM OOS P1 template.docx"
    if not os.path.exists(target_docx):
        target_docx = "EM OOS P1 template 0.docx"

    if os.path.exists(target_docx):
        try:
            from docxtpl import DocxTemplate
            doc = DocxTemplate(target_docx)
            doc.render(ctx)
            doc.save(docx_buf)
            docx_buf.seek(0)
        except Exception as e:
            st.error(f"Error rendering Word template: {e}")
            docx_buf = None
    else:
        docx_buf = None

    # 2. Generate Complete 7-Page PDF
    target_pdf = "EM OOS P1 template.pdf"
    if os.path.exists(target_pdf):
        try:
            from pypdf import PdfWriter, PdfReader
            
            # Map 157 Form 3.100.019.F01 fields
            pdf_map = {
                'Text Field57': ctx.get('oos_id', ''),
                'Text Field0': ctx.get('analyst_signature', ''),
                'Date Field0': ctx.get('test_date', ''),
                'Date Field1': ctx.get('date_initiated', ''),
                'Date Field2': ctx.get('date_of_incident', ''),
                'Date Field3': ctx.get('date_initiated', ''),
                'Text Field1': "Environmental Monitoring",
                'Text Field2': ctx['event_number'],
                'Text Field3': ctx['analyst_personnel_block'],
                'Text Field4': ctx['sample_name'],
                'Text Field5': "Plate",
                'Text Field6': ctx['lot_number'],
                'Text Field7': ctx['incident_description'],
                'Text Field8': "MICRO-SOP-2",
                'Text Field9': "23-Jul-2026",
                'Text Field10': "16",
                'Text Field11': ctx['action_level'].replace("≥", ">="),
                'Text Field12': ctx.get('manager_notified', 'Kathan Parikh'),
                'Text Field13': ctx['smart_comment_interview'],
                'Text Field14': "Not applicable",
                'Text Field15': "Yes, as per MICRO-SOP-2",
                'Text Field16': "Yes, as per MICRO-SOP-2",
                'Text Field17': ctx['smart_comment_records'],
                'Text Field18': "Yes, the analysts are trained and qualified by quality to perform the test",
                'Text Field19': "Not Applicable",
                'Text Field20': "Not Applicable",
                'Text Field21': "Yes, as per MICRO-SOP-2",
                'Text Field22': ctx.get('reagent_lot', "1011834770"),
                'Text Field23': ctx.get('reagent_exp', "25 Sep 2026"),
                'Text Field24': "Not Applicable",
                'Text Field25': "Not Applicable",
                'Text Field26': "Not Applicable",
                'Text Field27': "Not Applicable",
                'Text Field28': "Not Applicable",
                'Text Field29': "Not Applicable",
                'Text Field30': "Please see below",
                'Text Field31': "Please see below",
                'Text Field32': ctx.get('cr_display', "CR115 (E001737)"),
                'Text Field33': ctx.get('cr_exp', "Dec 2026"),
                'Text Field34': "N/A",
                'Text Field35': "N/A",
                'Text Field36': "Not Applicable",
                'Text Field37': "Not Applicable",
                'Text Field38': "Not Applicable",
                'Text Field39': "Not Applicable",
                'Text Field40': "Not Applicable",
                'Text Field41': "Not Applicable",
                'Text Field42': "Not Applicable",
                'Text Field45': "Not Applicable",
                'Text Field46': "Not Applicable",
                'Text Field47': "Not Applicable",
                'Text Field48': ctx.get('section_c_other', f"N/A QYC {datetime.now().strftime('%d-%b-%Y')}"),
                'Text Field49': interview_block,
                'Text Field50': records_block,
                'Text Field51': summary_block,
                'Text Field52': "",
                'Text Field53': ctx.get('writer_name', "Dhvanir Kansara"),
                'Text Field54': ctx.get('manager_signer', "Robin Seymour")
            }

            # Checkbox Yes/No defaults matching production PDF QA standards (EM is internal facility testing)
            yes_boxes = {4, 9, 10, 13, 16, 19, 24, 27, 28, 33, 36, 39, 42, 43, 48, 51, 52, 55, 60, 63, 66, 69, 72, 73, 78, 79, 87}
            for i in range(100):
                if i in yes_boxes:
                    pdf_map[f'Check Box{i}'] = '/Yes'
                else:
                    pdf_map[f'Check Box{i}'] = ''

            # Fill Form 1-6
            writer = PdfWriter(clone_from=target_pdf)
            for page in writer.pages:
                writer.update_page_form_field_values(page, pdf_map)
                
            # Generate Page 7 Attachment Table
            page7_pdf_buf = generate_em_tables_page_pdf(ctx)
            p7_reader = PdfReader(page7_pdf_buf)
            writer.add_page(p7_reader.pages[0])

            writer.write(pdf_buf)
            pdf_buf.seek(0)
        except Exception as e:
            st.error(f"Error rendering PDF template: {e}")
            pdf_buf = None
    else:
        pdf_buf = None

    return docx_buf, pdf_buf

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, top='single', bottom='single', left='single', right='single', color='auto', sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="{color}"/></w:tcBorders>')
    tcPr.append(tcBorders)

def format_date_em_tbl(d_str):
    if not d_str: return ""
    clean = re.sub(r'[\s\-]', '', str(d_str).strip())
    for fmt in ["%d%b%Y", "%d%b%y", "%Y%m%d"]:
        try:
            dt = datetime.strptime(clean, fmt)
            return dt.strftime("%d%b %Y").upper()
        except: pass
    return str(d_str).upper()

def generate_em_standalone_table_docx(ctx=None):
    """
    Generates a 100% faithful, plain, authentic EM Tables Word document (.docx)
    containing both Table 1 (Read Dates & Incubation Observation) and
    Table 2 (EM Bracketing Table, 18 rows) matching the exact style,
    font (Times New Roman 6-7pt), shading (#E8E8E8 / #D9D9D9 / #FFFFFF),
    and layout of 'EM table OOS-261187 11MAY2026.docx'.
    """
    if ctx is None:
        ctx = build_em_context()
        
    doc = docx.Document()
    
    # 1. Page Setup - Portrait 8.5 x 11.0 in, 1.0 in margins
    sec = doc.sections[0]
    sec.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    
    bsc_num = ctx.get('bsc_num', '1309')
    suite_num = ctx.get('suite_num', '117')
    analyst_init = ctx.get('analyst_initial', 'CGS')
    before_d = format_date_em_tbl(ctx.get('before_date', '08MAY 2026'))
    test_d = format_date_em_tbl(ctx.get('test_date', '11MAY 2026'))
    after_d = format_date_em_tbl(ctx.get('after_date', '12MAY 2026'))
    cfu_cnt = ctx.get('cfu_count', '1')
    etx_num = ctx.get('event_number', 'ETX-260518-0273')
    
    raw_org = ctx.get('manual_org', 'colony-like artifact (White, shiny, hardened, smooth area embeded on agar)')
    clean_org = raw_org.replace('colony-like artifact (', '').replace(')', '').replace('N/A', '').strip()
    if clean_org:
        org_display = f"N/A {clean_org}"
    else:
        org_display = "N/A"
        
    # --- TABLE 1: Read Dates and Incubation Observation ---
    t1_headers = [
        "Sampling Location",
        "Read Date\n(30-35°C, NLT 48h)",
        "Read\nBy",
        "CFU Count /\nObservation",
        "Read Date\n(20-25°C, NLT 5d)",
        "Read\nBy",
        "CFU Count /\nObservation",
        "Microbial Identification"
    ]
    t1_col_widths = [1.5, 0.7, 0.4, 0.9, 0.7, 0.4, 0.9, 1.0] # Sum = 6.5 inches
    
    t1_data = [
        ctx.get('sampling_location', f"{ctx.get('sampling_type', 'Surface Sampling')} Plate ({ctx.get('bsc_id', 'BSC E001309')})"),
        ctx.get('d_48h', '13 May 2026'),
        ctx.get('reader_48h', 'MC'),
        ctx.get('cfu_obs_48h', 'No microbial growth was observed'),
        ctx.get('d_5d', '18 May 2026'),
        ctx.get('reader_5d', 'SAS'),
        ctx.get('cfu_obs_5d', f"{cfu_cnt} CFU on {ctx.get('sample_name', '')}"),
        clean_org if clean_org else ctx.get('microbial_id', 'colony-like artifact')
    ]
    
    table1 = doc.add_table(rows=2, cols=8)
    table1.alignment = WD_TABLE_ALIGNMENT.LEFT
    table1.autofit = False
    
    # Table 1 Header Row
    for c_idx, text in enumerate(t1_headers):
        cell = table1.cell(0, c_idx)
        cell.width = Inches(t1_col_widths[c_idx])
        set_cell_background(cell, "E8E8E8")
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(6.0)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        
    # Table 1 Data Row
    for c_idx, text in enumerate(t1_data):
        cell = table1.cell(1, c_idx)
        cell.width = Inches(t1_col_widths[c_idx])
        set_cell_background(cell, "FFFFFF")
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(text))
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.0)
        r.bold = False
        r.font.color.rgb = RGBColor(0, 0, 0)
        
    # Spacer paragraph
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(6)
    p_spacer.paragraph_format.space_after = Pt(6)
    
    # --- TABLE 2: EM Bracketing Table (18 Rows) ---
    t2_rows_data = [
        # Row 0: Headers
        ["Environmental Monitoring (EM) Sampling Site", "Frequency", "Date\n(DDMM\nYYYY)", "Analyst (Initials)", "Day /Week(s)", "Observation*", "Environmental Monitoring Plate ETX ID", "Microbial ID", "Notes"],
        # Row 1: Section
        ["Personnel EM Bracketing"],
        # Row 2-4: Pers
        ["Personal (Left Touch and Right Touch)", "Daily", before_d, analyst_init, "Date Before Testing", "No growth", "N/A", "N/A", "None"],
        ["Personal (Left Touch and Right Touch)", "Daily", test_d, analyst_init, "Date of Testing", "No growth", "N/A", "N/A", "None"],
        ["Personal (Left Touch and Right Touch)", "Daily", after_d, analyst_init, "Date After Testing", "No growth", "N/A", "N/A", "None"],
        # Row 5: Section BSC
        [f"Biological Safety Cabinet EM Bracketing Biological Safety Cabinet (BSC) {bsc_num}"],
        # Row 6-8: BSC Surface
        ["Surface Sampling of ISO 5 (4 locations)", "Daily", "09MAY 2026", "VV", "Date Before Testing", "No growth", "", "", "None"],
        ["Surface Sampling of ISO 5\n(4 locations)", "Daily", test_d, f"{analyst_init}, DT, GA, VV ", "Date of Testing", f"{cfu_cnt} CFU on S1 for {analyst_init}", etx_num, org_display, "None"],
        ["Surface Sampling of ISO 5 (4 locations)", "Daily", after_d, "PG, DT, HS, ELB, GA, VV", "Date After Testing", "No growth", "N/A", "N/A", "None"],
        # Row 9-11: BSC Settling
        ["Settling Sampling of ISO 5 (2 locations)", "Daily", "09MAY 2026", "VV", "Date Before Testing", "No growth", "N/A", "N/A", "None"],
        ["Settling Sampling of ISO 5 (2 locations)", "Daily", test_d, f"{analyst_init}, DT, GA, VV ", "Date of Testing", "No growth", "N/A", "N/A", "None"],
        ["Settling Sampling of ISO 5 (2 locations)", "Daily", after_d, "PG, DT, HS, ELB, GA, VV", "Date After Testing", "No growth", "N/A", "N/A", "None"],
        # Row 12: Section Air
        [f"Weekly Active Air Sampling Bracketing {suite_num}"],
        # Row 13-14: Air
        ["Active Air Sampling of Cleanrooms", "Weekly", "29APR 2026", "SMO", "Week (On of Testing Date)", "No growth", "N/A", "N/A", "None"],
        ["Active Air Sampling of Cleanrooms", "Weekly", "07MAY 2026", "SMO", "Week (On or After Testing Date)", f"2 CFU for {suite_num} Air", "ETX-260518-0263", "Gram (+) cocci\n Gram (+) pleomorphic rods", "None"],
        # Row 15: Section Surf
        [f"Surface Sampling of Anteroom and Cleanroom Bracketing {suite_num}"],
        # Row 16-17: Cleanroom Surf
        ["Surface Sampling of Cleanrooms", "Weekly", "29APR 2026", "SMO", "Week (On of Testing Date)", "No growth", "N/A", "N/A", "None"],
        ["Surface Sampling of Cleanrooms", "Weekly", "07MAY 2026", "SMO", "Week (On or After Testing Date)", "No growth", "N/A", "N/A", "None"]
    ]
    
    col_widths = [1.1125, 0.5271, 0.5660, 0.6243, 0.7611, 0.9368, 0.9014, 0.7201, 0.4222]
    
    table2 = doc.add_table(rows=len(t2_rows_data), cols=9)
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT
    table2.autofit = False
    
    for r_idx, row_data in enumerate(t2_rows_data):
        row = table2.rows[r_idx]
        is_hdr = (r_idx == 0)
        is_sec = (r_idx in [1, 5, 12, 15])
        
        if is_sec:
            first_c = row.cells[0]
            for c in row.cells[1:]:
                first_c.merge(c)
            set_cell_background(first_c, "D9D9D9")
            set_cell_borders(first_c)
            p = first_c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(row_data[0])
            r.font.name = "Times New Roman"
            r.font.size = Pt(7.0)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)
            continue
            
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            set_cell_borders(cell)
            
            if is_hdr:
                set_cell_background(cell, "E8E8E8")
            else:
                set_cell_background(cell, "FFFFFF")
                
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            r.font.name = "Times New Roman"
            r.font.size = Pt(6.0) if is_hdr else Pt(7.0)
            r.bold = is_hdr
            r.font.color.rgb = RGBColor(0, 0, 0)
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
