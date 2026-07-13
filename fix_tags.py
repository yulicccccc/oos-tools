import docx

def fix_tags(filepath):
    doc = docx.Document(filepath)
    for p in doc.paragraphs:
        if 'subculture _name' in p.text:
            p.text = p.text.replace('subculture _name', 'subculture_name')
        if 'subculture _initial' in p.text:
            p.text = p.text.replace('subculture _initial', 'subculture_initial')
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if 'subculture _name' in p.text:
                        p.text = p.text.replace('subculture _name', 'subculture_name')
                    if 'subculture _initial' in p.text:
                        p.text = p.text.replace('subculture _initial', 'subculture_initial')
    doc.save(filepath)

if __name__ == "__main__":
    fix_tags('USP71 OOS P1 template 0.docx')
